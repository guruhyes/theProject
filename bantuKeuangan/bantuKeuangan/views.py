from django.shortcuts import render
#print modul
from django.core.files.storage import FileSystemStorage
from django.http import HttpResponse
import io
from django.template.loader import render_to_string
from weasyprint import HTML
from datetime import datetime
from datetime import timedelta
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.worksheet.table import Table, TableStyleInfo
from docx import Document
from docxtpl import DocxTemplate
import win32com.client
#Import Model Siswa
# Create your views here.
def index(request):
    return render(request,'base.html')
def print(request):
    # AllSiswa=SiswaModel.objects.all()
    # context = {
    #     "AllSiswa":AllSiswa
    # }

    html_string = render_to_string('weasyPrint.html')
    html = HTML(string=html_string, base_url=request.build_absolute_uri())
    html.write_pdf(target='\\tmp\\mypdf.pdf')
    fs = FileSystemStorage('/tmp')
    with fs.open('mypdf.pdf') as pdf:
        response = HttpResponse(pdf, content_type='application/pdf')
        response['Content-Disposition'] = 'filename=Report.pdf'
        return response
    return response
    # pass

def exportExcel(request):

    """
    Downloads all movies as Excel file with a single worksheet
    """
    #movie_queryset = Movie.objects.all()
    
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    )
    response['Content-Disposition'] = 'attachment; filename={date}-movies.xlsx'.format(
        date=datetime.now().strftime('%Y-%m-%d'),
    )
    workbook = Workbook()
    
    # Get active worksheet/tab
    worksheet = workbook.active
    worksheet.title = 'Movies'

    # Define the titles for columns
    columns = [
        'ID',
        'Title',
        'Description',
        'Length',
        'Rating',
        'Price',
    ]
    row_num = 1

    # Assign the titles for each cell of the header
    for col_num, column_title in enumerate(columns, 1):
        cell = worksheet.cell(row=row_num, column=col_num)
        cell.value = column_title

    # # Iterate through all movies
    # for movie in movie_queryset:
    #     row_num += 1
        
    #     # Define the data for each cell in the row 
    #     row = [
    #         movie.pk,
    #         movie.title,
    #         movie.description,
    #         movie.length_in_minutes,
    #         movie.rating,
    #         movie.price,
    #     ]
        
    #     # Assign the data for each cell of the row 
    #     for col_num, cell_value in enumerate(row, 1):
    #         cell = worksheet.cell(row=row_num, column=col_num)
    #         cell.value = cell_value
    for x in range(5):
        cell = worksheet.cell(row=2, column=x+1)
        cell.value = 3

    f = worksheet.cell(row=2, column=6)
    f.value = '=sum({3}{2}:{0}{1})'.format('E', 2, 2,'A')
    f.font = Font(name='Calibri', bold=True)
    worksheet.merge_cells(start_row=1, start_column=6, end_row=1, end_column=8)

    tab = Table(displayName="Table1", ref="A1:E10")

    # Add a default style with striped rows and banded columns
    style = TableStyleInfo(name="TableStyleMedium9", showFirstColumn=False,
                        showLastColumn=False, showRowStripes=True, showColumnStripes=True)
    tab.tableStyleInfo = style

    '''
    Table must be added using ws.add_table() method to avoid duplicate names.
    Using this method ensures table name is unque through out defined names and all other table name. 
    '''
    worksheet.add_table(tab)
    workbook.save(response)

    return response
def doc(request):
    return render(request,'doc.html')
def exportWord(requet):
    word = win32com.client.Dispatch('Word.Application')

    doc = word.Documents.Add('example.html')
    doc.SaveAs('example.doc', FileFormat=10)
    doc.Close()

    word.Quit()
def replaceWord(request):
    doc=Document('./need/test.docx')
    Dictionary = {"sea": "ocean", "find_this_text":"new_text"}
    for i in Dictionary:
        for p in doc.paragraphs:
            if p.text.find(i)>=0:
                p.text=p.text.replace(i,Dictionary[i])
    #save changed document
    doc.save('./need/test2.docx')
def doc_test(request):
    doc = DocxTemplate("./need/template.docx")
    # ... your other code ...
    context = { 'sea' : "World company",'find_this_text':"new text" }
    doc.render(context)
    doc.save("./need/generated_doc.docx")
    doc_io = io.BytesIO() # create a file-like object
    doc.save(doc_io) # save data to file-like object
    doc_io.seek(0) # go to the beginning of the file-like object

    response = HttpResponse(doc_io.read())

    # Content-Disposition header makes a file downloadable
    response["Content-Disposition"] = "attachment; filename=generated_doc.docx"

    # Set the appropriate Content-Type for docx file
    response["Content-Type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    return response